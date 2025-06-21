import tkinter as tk
from tkinter import filedialog, messagebox
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from PIL import Image
import csv
import os
from docx import Document
from docx.shared import Inches

def get_frame(parent):
    frame = tk.Frame(parent, bg="white")

    tk.Label(frame, text="qPCR Reaction Planner", font=("Segoe UI", 14, "bold"), bg="white").pack(pady=10)

    reagents = ["MasterMix", "Forward Primer", "Reverse Primer", "Probe", "Template", "Water"]
    entries = {}

    tk.Label(frame, text="Volume per Reaction (µL):", font=("Segoe UI", 10, "underline"), bg="white").pack()

    for r in reagents:
        row = tk.Frame(frame, bg="white")
        row.pack()
        tk.Label(row, text=f"{r}:", width=20, anchor="e", bg="white").pack(side="left")
        e = tk.Entry(row, width=12)
        e.pack(side="left")
        entries[r] = e

    tk.Label(frame, text="Number of Reactions:", bg="white").pack(pady=(10, 0))
    reps_entry = tk.Entry(frame)
    reps_entry.pack()

    tk.Label(frame, text="Overage % [Default = 10]:", bg="white").pack(pady=(5, 0))
    overage_entry = tk.Entry(frame)
    overage_entry.pack()

    result = tk.StringVar()
    result_label = tk.Label(frame, textvariable=result, bg="white", fg="green", justify="left")
    result_label.pack(pady=(10, 0))

    canvas_frame = tk.Frame(frame, bg="white")
    canvas_frame.pack(pady=10)

    results_cache = {}

    def calculate():
        try:
            reps = int(reps_entry.get())
            over = float(overage_entry.get()) if overage_entry.get() else 10
            multiplier = 1 + (over / 100)
            results_cache.clear()

            for r in reagents:
                val = entries[r].get().strip()
                if not val:
                    raise ValueError(f"{r} volume missing")
                total = round(float(val) * reps * multiplier, 2)
                results_cache[r] = total

            total_sum = round(sum(results_cache.values()), 2)
            lines = [f"{r}: {v} µL" for r, v in results_cache.items()]
            lines.append(f"Total: {total_sum} µL")
            result.set("\n".join(lines))

            for widget in canvas_frame.winfo_children():
                widget.destroy()

            fig, ax = plt.subplots(figsize=(6, 3))
            ax.bar(results_cache.keys(), results_cache.values(), color="#4a90e2")
            ax.set_ylabel("Volume (µL)")
            ax.set_title("qPCR Reagent Volumes")
            ax.set_xticklabels(results_cache.keys(), rotation=45, ha="right")

            fig.tight_layout()
            fig.savefig("qpcr_plot_temp.png", dpi=150)  # temp image for docx export

            canvas = FigureCanvasTkAgg(fig, master=canvas_frame)
            canvas.draw()
            canvas.get_tk_widget().pack()

        except Exception as e:
            result.set(f"Error: {e}")

    def export_docx():
        if not results_cache:
            messagebox.showinfo("Export Error", "Please calculate first.")
            return

        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            title="Export qPCR Output"
        )
        if not path:
            return

        try:
            doc = Document()
            doc.add_heading("qPCR Reaction Report", level=1)
            doc.add_paragraph(f"Number of Reactions: {reps_entry.get()}")
            doc.add_paragraph(f"Overage: {overage_entry.get() or 10}%")

            doc.add_paragraph("Total Reagent Volumes:")
            table = doc.add_table(rows=1, cols=2)
            hdr = table.rows[0].cells
            hdr[0].text = "Reagent"
            hdr[1].text = "Volume (µL)"

            for r, v in results_cache.items():
                row = table.add_row().cells
                row[0].text = r
                row[1].text = str(v)

            table.add_row().cells[0].text = "Total"
            table.rows[-1].cells[1].text = str(round(sum(results_cache.values()), 2))

            if os.path.exists("qpcr_plot_temp.png"):
                doc.add_picture("qpcr_plot_temp.png", width=Inches(5.5))
            doc.save(path)
            messagebox.showinfo("Export Complete", f"Exported to {os.path.basename(path)}")
        except Exception as e:
            messagebox.showerror("Export Failed", str(e))

    tk.Button(frame, text="Calculate & Plot", command=calculate, bg="#4a90e2", fg="white").pack(pady=6)
    tk.Button(frame, text="Export as Word (.docx)", command=export_docx, bg="#2b8a3e", fg="white").pack(pady=3)

    return frame